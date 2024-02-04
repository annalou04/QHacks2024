'''
converting an audio file to an ai generated image
author: Shreya M
'''


import openai
from openai import OpenAI
client = OpenAI()
from docx import Document
import requests
from io import BytesIO
from docx.shared import Inches


def transcribe_audio(audio_file_path):
	"""Take in an audio file and transcribes it (Whisper model)"""  
	audio_file= open(audio_file_path, "rb")
	transcript = client.audio.transcriptions.create(
		model="whisper-1", 
		file=audio_file
	)
	print(transcript)
	stripped_transcript = str(transcript)[20:-2]
	return(stripped_transcript)

def save_as_docx(transcription, filename):
	"""Converts the raw text to a Word document"""
	doc = Document()
	doc.add_paragraph(str(transcription))
	doc.save(filename)

def generate_image(prompt):
    response = client.images.generate(
        model="dall-e-2",  # Use the appropriate DALL-E model
        prompt=prompt,
        size="1024x1024",
        quality="standard",
        n=1,
    )

    image_url = response.data[0].url
    return image_url

def save_image_to_word(image_url, output_path="outputtxt.docx"):
    # Download the image
    image_response = requests.get(image_url)
    image_data = BytesIO(image_response.content)

    # Create a Word document
    doc = Document()
    doc.add_heading('Generated Image', level=1)

    # Add the image to the document
    doc.add_picture(image_data, width=Inches(4))

    # Save the document
    doc.save(output_path)

    print(f"Word document saved at {output_path}")
#computer specifc pathname
audio_file_path = "/Users/Shreds/Desktop/qhacks/harvard.wav"
transcription = transcribe_audio(audio_file_path)
image_url = generate_image(transcription)  # Obtain the image URL
save_as_docx(transcription, 'outputtxt.docx')
save_image_to_word(image_url, 'outputtxt.docx')  # Pass image_url to the function


