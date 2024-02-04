'''
converting a text file into an AI generated image
author:Shreya M

'''
from openai import OpenAI
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches

def generate_image(prompt):
    client = OpenAI()

    response = client.images.generate(
        model="dall-e-2",
        prompt=prompt,
        size="1024x1024",
        quality="standard",
        n=1,
    )

    image_url = response.data[0].url
    return image_url

def save_image_to_word(prompt, output_path="output.docx"):
    image_url = generate_image(prompt)
    image_response = requests.get(image_url)
    image_data = BytesIO(image_response.content)
    doc = Document()
    doc.add_heading('Generated Image', level=1)
    
    doc.add_picture(image_data, width=Inches(4))

    doc.save(output_path)

    print(f"Word document saved at {output_path}")


def readingFile(filePath):
    f = open(filePath,'r')
    prompt = f.read().strip()
    f.close()
    return prompt
#computer specifc pathname
filePath="/Users/Shreds/Desktop/qhacks/test.txt"
prompt = readingFile(filePath)
save_image_to_word(prompt)
    
    
